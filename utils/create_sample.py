import docx
import os

def create_sample_docx(filename="sample_contract.docx"):
    doc = docx.Document()
    doc.add_heading('Service Agreement', 0)
    
    doc.add_paragraph('This Service Agreement ("Agreement") is entered into by and between Company A ("Provider") and Client B ("Client").')
    
    doc.add_heading('1. Services', level=1)
    doc.add_paragraph('Provider agrees to deliver software development services including but not limited to Web Development, API Integration, and Cloud Deployment.')
    
    doc.add_heading('2. Payment Terms', level=1)
    doc.add_paragraph('Client shall pay Provider at a rate of $150 per hour. Invoices will be issued monthly and are due within 30 days of receipt.')
    
    doc.add_heading('3. Confidentiality', level=1)
    doc.add_paragraph('Both parties agree to maintain the confidentiality of all proprietary information disclosed during the term of this Agreement.')
    
    doc.add_heading('4. Termination', level=1)
    doc.add_paragraph('Either party may terminate this Agreement with 14 days written notice.')
    
    doc.save(filename)
    print(f"Created {filename}")

if __name__ == "__main__":
    create_sample_docx()
